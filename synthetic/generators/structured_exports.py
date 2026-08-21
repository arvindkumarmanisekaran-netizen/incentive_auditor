from __future__ import annotations

import csv
import json
import random
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook

SEED = 42

FORMATS = [
    "csv",
    "json",
    "xlsx",
    "docx",
]


def export_structured_data(
    documents: dict[str, list[dict[str, Any]]],
    output_dir: Path,
) -> None:

    output_dir.mkdir(
        parents=True,
        exist_ok=True,
    )

    rng = random.Random(SEED)

    for document_name, records in documents.items():

        if not records:
            continue

        shuffled_records = [dict(record) for record in records]

        rng.shuffle(shuffled_records)

        format_records = split_records(
            shuffled_records,
            len(FORMATS),
        )

        base_name = Path(document_name).stem

        for file_format, rows in zip(
            FORMATS,
            format_records,
        ):

            if not rows:
                continue

            filename = f"{base_name}.{file_format}"

            output_path = output_dir / filename

            if file_format == "csv":
                export_csv(
                    rows,
                    output_path,
                )

            elif file_format == "json":
                export_json(
                    rows,
                    output_path,
                )

            elif file_format == "xlsx":
                export_xlsx(
                    rows,
                    output_path,
                )

            elif file_format == "docx":
                export_docx(
                    rows,
                    output_path,
                )

            print(f"  {filename:40} " f"{len(rows):,} records")


# ============================================================
# SPLIT DATA
# ============================================================


def split_records(
    records: list[dict[str, Any]],
    parts: int,
) -> list[list[dict[str, Any]]]:

    result: list[list[dict[str, Any]]] = [[] for _ in range(parts)]

    # Round-robin distribution ensures different records
    # go into each format while keeping sizes balanced.
    for index, record in enumerate(records):

        result[index % parts].append(record)

    return result


# ============================================================
# CSV
# ============================================================


def export_csv(
    records: list[dict[str, Any]],
    path: Path,
) -> None:

    if not records:
        return

    columns = collect_columns(
        records,
    )

    with open(
        path,
        "w",
        newline="",
        encoding="utf-8",
    ) as file:

        writer = csv.DictWriter(
            file,
            fieldnames=columns,
            extrasaction="ignore",
        )

        writer.writeheader()

        for record in records:

            writer.writerow(
                normalize_record(
                    record,
                    columns,
                )
            )


# ============================================================
# JSON
# ============================================================


def export_json(
    records: list[dict[str, Any]],
    path: Path,
) -> None:

    with open(
        path,
        "w",
        encoding="utf-8",
    ) as file:

        json.dump(
            records,
            file,
            indent=2,
            ensure_ascii=False,
            default=str,
        )


# ============================================================
# XLSX
# ============================================================


def export_xlsx(
    records: list[dict[str, Any]],
    path: Path,
) -> None:

    if not records:
        return

    columns = collect_columns(
        records,
    )

    workbook = Workbook()

    worksheet = workbook.active

    worksheet.title = "Data"

    worksheet.append(columns)

    for record in records:

        normalized = normalize_record(
            record,
            columns,
        )

        worksheet.append([normalize_cell_value(normalized[column]) for column in columns])

    workbook.save(path)


# ============================================================
# DOCX
# ============================================================


def export_docx(
    records: list[dict[str, Any]],
    path: Path,
) -> None:

    if not records:
        return

    columns = collect_columns(
        records,
    )

    document = Document()

    document.add_heading(
        "Structured Data Export",
        level=1,
    )

    document.add_paragraph(f"Total records: {len(records)}")

    table = document.add_table(
        rows=1,
        cols=len(columns),
    )

    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    for index, column in enumerate(columns):

        header_cells[index].text = str(column)

    for record in records:

        normalized = normalize_record(
            record,
            columns,
        )

        row_cells = table.add_row().cells

        for index, column in enumerate(columns):

            value = normalized.get(column)

            row_cells[index].text = format_docx_value(value)

    document.save(path)


# ============================================================
# HELPERS
# ============================================================


def collect_columns(
    records: list[dict[str, Any]],
) -> list[str]:

    columns: list[str] = []

    seen: set[str] = set()

    for record in records:

        for key in record.keys():

            if key not in seen:

                seen.add(key)

                columns.append(key)

    return columns


def normalize_record(
    record: dict[str, Any],
    columns: list[str],
) -> dict[str, Any]:

    return {column: record.get(column) for column in columns}


def normalize_cell_value(
    value: Any,
) -> Any:

    if value is None:
        return ""

    if isinstance(
        value,
        (
            dict,
            list,
            tuple,
            set,
        ),
    ):

        return json.dumps(
            value,
            ensure_ascii=False,
            default=str,
        )

    return value


def format_docx_value(
    value: Any,
) -> str:

    if value is None:
        return ""

    if isinstance(
        value,
        (
            dict,
            list,
            tuple,
            set,
        ),
    ):

        return json.dumps(
            value,
            ensure_ascii=False,
            default=str,
        )

    return str(value)
