from __future__ import annotations

import csv
import io
import json
import logging
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document

SUPPORTED_EXTENSIONS = {
    "csv",
    "xlsx",
    "json",
    "docx",
}


logger = logging.getLogger(__name__)


def get_extension(
    filename: str,
) -> str:
    return Path(filename).suffix.lower().lstrip(".")


def validate_extension(
    filename: str,
) -> str:
    extension = get_extension(filename)

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            "Unsupported document type. " "Supported formats are CSV, XLSX, JSON, and DOCX."
        )

    return extension


def normalize_record(
    record: dict[str, Any],
) -> dict[str, Any]:
    """
    Normalize parser output so downstream
    services receive consistent Python values.
    """

    normalized: dict[str, Any] = {}

    for key, value in record.items():
        column_name = str(key).strip()

        if not column_name:
            continue

        if value is None:
            normalized[column_name] = None
            continue

        try:
            if pd.isna(value):
                normalized[column_name] = None
                continue
        except (TypeError, ValueError):
            pass

        # Pandas Timestamp -> ISO date/time
        if isinstance(
            value,
            pd.Timestamp,
        ):
            normalized[column_name] = value.isoformat()
            continue

        # Convert numpy scalar values
        # into normal Python values.
        if hasattr(
            value,
            "item",
        ):
            try:
                value = value.item()

            except (
                ValueError,
                AttributeError,
            ):
                pass

        # Normalize string whitespace.
        if isinstance(value, str):
            value = value.strip()

            if value == "":
                value = None

        normalized[column_name] = value

    return normalized


def parse_csv(
    content: bytes,
) -> list[dict[str, Any]]:
    text = content.decode("utf-8-sig")

    buffer = io.StringIO(text)

    reader = csv.DictReader(buffer)

    if not reader.fieldnames:
        raise ValueError("CSV file does not contain a header row.")

    records = [normalize_record(dict(row)) for row in reader]

    return records


def parse_xlsx(
    content: bytes,
) -> list[dict[str, Any]]:
    buffer = io.BytesIO(content)

    dataframe = pd.read_excel(
        buffer,
        engine="openpyxl",
    )

    if dataframe.empty:
        return []

    records = dataframe.to_dict(
        orient="records",
    )

    return [normalize_record(record) for record in records]


def parse_json(
    content: bytes,
) -> list[dict[str, Any]]:
    decoded = content.decode(
        "utf-8-sig",
    )

    parsed = json.loads(decoded)

    # ----------------------------------------
    # JSON array
    #
    # [
    #   {...},
    #   {...}
    # ]
    # ----------------------------------------

    if isinstance(
        parsed,
        list,
    ):
        if not all(
            isinstance(
                item,
                dict,
            )
            for item in parsed
        ):
            raise ValueError("JSON arrays must contain objects.")

        return [normalize_record(item) for item in parsed]

    # ----------------------------------------
    # JSON object
    #
    # {
    #   "records": [...]
    # }
    #
    # or
    #
    # {
    #   "sales": [...]
    # }
    # ----------------------------------------

    if isinstance(
        parsed,
        dict,
    ):
        for value in parsed.values():
            if isinstance(
                value,
                list,
            ) and all(
                isinstance(
                    item,
                    dict,
                )
                for item in value
            ):
                return [normalize_record(item) for item in value]

        # Single-record JSON object
        return [normalize_record(parsed)]

    raise ValueError("Unsupported JSON structure.")


def parse_docx(
    content: bytes,
) -> list[dict[str, Any]]:
    """
    Parse structured DOCX files.

    Expected format:

    First table row:
        column headers

    Remaining rows:
        record values

    Multiple tables are supported.
    """

    try:
        document = Document(io.BytesIO(content))

    except Exception as exc:
        raise ValueError(f"Invalid DOCX document: {exc}") from exc

    if not document.tables:
        raise ValueError("DOCX document does not contain " "a structured table.")

    records: list[dict[str, Any]] = []

    for table in document.tables:
        if not table.rows:
            continue

        # ----------------------------------------
        # Header row
        # ----------------------------------------

        headers = [cell.text.strip() for cell in table.rows[0].cells]

        if not any(headers):
            continue

        # ----------------------------------------
        # Data rows
        # ----------------------------------------

        for row in table.rows[1:]:
            values = [cell.text.strip() for cell in row.cells]

            # Skip blank rows
            if not any(values):
                continue

            record: dict[str, Any] = {}

            for index, header in enumerate(headers):
                if not header:
                    continue

                value = values[index] if index < len(values) else None

                record[header] = value if value not in ("", None) else None

            normalized = normalize_record(record)

            if normalized:
                records.append(normalized)

    if not records:
        raise ValueError("DOCX document contains no records.")

    return records


def parse_document(
    filename: str,
    content: bytes,
) -> list[dict[str, Any]]:
    """
    Main entry point used by the document
    processing service.
    """

    extension = validate_extension(filename)

    if not content:
        raise ValueError("Uploaded document is empty.")

    if extension == "csv":
        records = parse_csv(content)

    elif extension == "xlsx":
        records = parse_xlsx(content)

    elif extension == "json":
        records = parse_json(content)

    elif extension == "docx":
        records = parse_docx(content)

    else:
        raise ValueError(f"No parser configured for {extension}")

    if not records:
        raise ValueError("Document contains no records.")

    return records


def extract_source_columns(
    records: list[dict[str, Any]],
) -> list[str]:
    """
    Collect every source column present
    in the document.

    Do not rely only on the first row because
    JSON records may have inconsistent keys.
    """

    columns: list[str] = []

    seen: set[str] = set()

    for record in records:
        for column in record.keys():
            if column not in seen:
                seen.add(column)

                columns.append(column)

    return columns
